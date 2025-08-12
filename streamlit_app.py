
import streamlit as st
from docx import Document
from datetime import date
import io

st.set_page_config(page_title="Gratuity Form I - Auto Fill", page_icon="🧾", layout="centered")

st.title("🧾 Gratuity Form ‘I’ (Auto-Fill)")
st.write("Enter details below and download a print-ready Word document.")

with st.form("form_i"):
    col1, col2 = st.columns(2)
    with col1:
        employee_name = st.text_input("Employee Name *")
        employee_address = st.text_area("Employee Address *")
        last_department = st.text_input("Last Department / Branch / Section")
        post_ticket_no = st.text_input("Post held / Ticket or Serial No.")
        date_of_appointment = st.date_input("Date of Appointment", value=date(2020,1,1))
        termination_date = st.date_input("Date of Termination", value=date.today())
    with col2:
        termination_cause = st.text_input("Cause of Termination (e.g., retirement, resignation)")
        total_service_period = st.text_input("Total Period of Service (e.g., 5 years 3 months)")
        last_wages = st.text_input("Amount of Wages Last Drawn")
        gratuity_amount = st.text_input("Amount of Gratuity Claimed")
        payment_mode = st.selectbox("Payment Mode", ["cash", "open or crossed bank cheque"])
        place = st.text_input("Place")

    establishment_full_address = st.text_area("Establishment (name/description with full address) *")
    disability_details = st.text_area("Disability Details (if applicable)")
    witness_details = st.text_area("Evidences/Witnesses (if applicable)")
    effective_date = st.date_input("Effective Date", value=date.today())
    signature_name = st.text_input("Applicant Signature Name")
    application_date = st.date_input("Application Date", value=date.today())

    submitted = st.form_submit_button("Generate Form")

if submitted:
    # Load template and replace placeholders
    template_path = "Form_I_template.docx"
    doc = Document(template_path)

    mapping = {
        "employee_name": employee_name,
        "employee_address": employee_address,
        "last_department": last_department,
        "post_ticket_no": post_ticket_no,
        "date_of_appointment": date_of_appointment.strftime("%Y-%m-%d"),
        "termination_date_and_cause": f"{termination_date.strftime('%Y-%m-%d')} - {termination_cause}",
        "total_service_period": total_service_period,
        "last_wages": last_wages,
        "gratuity_amount": gratuity_amount,
        "payment_mode": payment_mode,
        "place": place,
        "establishment_full_address": establishment_full_address,
        "disability_details": disability_details,
        "witness_details": witness_details,
        "effective_date": effective_date.strftime("%Y-%m-%d"),
        "signature_name": signature_name,
        "application_date": application_date.strftime("%Y-%m-%d"),
        "reason": termination_cause or "superannuation/retirement/resignation",
    }

    # Replace placeholders
    def replace_placeholders(doc, mapping):
        for p in doc.paragraphs:
            text = p.text
            changed = False
            for k, v in mapping.items():
                placeholder = "{{" + k + "}}"
                if placeholder in text:
                    text = text.replace(placeholder, str(v))
                    changed = True
            if changed:
                # Rebuild runs
                for i in range(len(p.runs)-1, -1, -1):
                    p._element.remove(p.runs[i]._element)
                p.add_run(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = p.text
                        changed = False
                        for k, v in mapping.items():
                            placeholder = "{{" + k + "}}"
                            if placeholder in text:
                                text = text.replace(placeholder, str(v))
                                changed = True
                        if changed:
                            for i in range(len(p.runs)-1, -1, -1):
                                p._element.remove(p.runs[i]._element)
                            p.add_run(text)

    replace_placeholders(doc, mapping)

    # Save to buffer for download
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    st.success("Your form is ready!")
    st.download_button(
        label="⬇️ Download Filled Form (DOCX)",
        data=buf,
        file_name=f"Gratuity_Form_I_{employee_name.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    st.info("Open the downloaded file and Print from Word. If you need a PDF, export to PDF from Word's Save As menu.")
