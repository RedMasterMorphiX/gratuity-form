
from docx import Document
import sys
from datetime import date

# Usage:
# python fill_form.py output.docx key=value key=value ...
# Example:
# python fill_form.py filled_form.docx employee_name="Rahul Sharma" establishment_full_address="ABC Pvt Ltd, Kolkata"

def load_template(path):
    return Document(path)

def replace_placeholders(doc, mapping):
    # Replace in paragraphs
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text or "{{" + k + "}}" in p.text:
                inline = p.runs
                # rebuild paragraph text to preserve runs
                new_text = p.text
                new_text = new_text.replace("{{" + k + "}}", str(v))
                new_text = new_text.replace(k, str(v))
                # Clear existing runs
                for i in range(len(inline)-1, -1, -1):
                    p._element.remove(inline[i]._element)
                p.add_run(new_text)

    # Replace in tables (if any added later)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in mapping.items():
                        if "{{" + k + "}}" in p.text or k in p.text:
                            inline = p.runs
                            new_text = p.text
                            new_text = new_text.replace("{{" + k + "}}", str(v))
                            new_text = new_text.replace(k, str(v))
                            for i in range(len(inline)-1, -1, -1):
                                p._element.remove(inline[i]._element)
                            p.add_run(new_text)

def parse_args(argv):
    # argv like: ["output.docx", 'employee_name="Rahul"']
    if len(argv) < 2:
        print("Usage: python fill_form.py output.docx key=value ...")
        sys.exit(1)
    out_path = argv[0]
    mapping = {}
    for part in argv[1:]:
        if "=" in part:
            k, v = part.split("=", 1)
            v = v.strip()
            if (v.startswith('"') and v.endswith('"')) or (v.startswith("'") and v.endswith("'")):
                v = v[1:-1]
            mapping[k.strip()] = v
    return out_path, mapping

if __name__ == "__main__":
    template_path = "Form_I_template.docx"
    out_path, mapping = parse_args(sys.argv[1:])

    # sensible defaults if not provided
    defaults = {
        "reason": "superannuation/retirement/resignation",
        "effective_date": str(date.today()),
        "disability_details": "",
        "witness_details": "",
        "payment_mode": "cash",
        "signature_name": "",
        "place": "",
        "application_date": str(date.today())
    }
    for k, v in defaults.items():
        mapping.setdefault(k, v)

    doc = load_template(template_path)
    replace_placeholders(doc, mapping)
    doc.save(out_path)
    print(f"Filled form saved to {out_path}")
